"""可扩展文档加载器：PDF / DOCX / CSV。"""
from dataclasses import dataclass, field
from typing import List, Dict
from abc import ABC, abstractmethod
from pathlib import Path


@dataclass
class Document:
    content: str
    metadata: Dict[str, str] = field(default_factory=dict)


class BaseLoader(ABC):
    """加载器基类，新增格式只需子类化。"""

    @abstractmethod
    def load(self, file_path: Path) -> List[Document]:
        ...


class PDFLoader(BaseLoader):
    def load(self, file_path: Path) -> List[Document]:
        import pdfplumber

        docs = []
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        content=text.strip(),
                        metadata={"source": file_path.name, "page": str(i + 1), "format": "pdf"},
                    ))
        return docs


class DOCXLoader(BaseLoader):
    def load(self, file_path: Path) -> List[Document]:
        from docx import Document as DocxDoc

        doc = DocxDoc(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return []
        return [Document(
            content="\n".join(paragraphs),
            metadata={"source": file_path.name, "format": "docx"},
        )]


class CSVLoader(BaseLoader):
    """CSV FAQ 加载器，将每行拼接为 question + answer 格式。"""

    def __init__(self, question_col: str = "问题", answer_col: str = "答案"):
        self.question_col = question_col
        self.answer_col = answer_col

    def load(self, file_path: Path) -> List[Document]:
        import csv

        docs = []
        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                q = row.get(self.question_col, "").strip()
                a = row.get(self.answer_col, "").strip()
                if q and a:
                    docs.append(Document(
                        content=f"Q: {q}\nA: {a}",
                        metadata={"source": file_path.name, "format": "csv"},
                    ))
        return docs


# 格式 → 加载器注册表
LOADER_REGISTRY: Dict[str, BaseLoader] = {
    ".pdf": PDFLoader(),
    ".docx": DOCXLoader(),
    ".csv": CSVLoader(),
}
